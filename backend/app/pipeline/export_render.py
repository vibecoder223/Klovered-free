"""Branded proposal export renderers — DOCX and PDF (port of lib/docx-export.ts
`renderDocx` and the inline `renderPdf` from app/api/exports/generate/route.ts).

Both produce a polished RFP-response document from the same data: a cover block
(kicker + title + prepared-for/submitted-by + date + accent rule), optional
proposal narrative sections, the Q&A body (optionally grouped per source
document), inline or footnote citations, and a page-numbered footer.

Typography:
- DOCX sets "Aptos" as the document font — the current Microsoft 365 default,
  so real Word/Docs users see the intended clean geometric-humanist sans; the
  rest gracefully substitute. Sizes/colors/spacing match the TS 1:1.
- PDF uses Helvetica (a reportlab built-in Type-1 face, embedded-free and
  therefore identical on every host — container, CI, laptop) with a deliberate
  type scale, accent headings, hairline rules, and justified body copy.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ---- shared palette (matches the app's --fg ramp / lib/docx-export.ts) ----
INK = "1A1A17"      # near-black body ink
MUTE = "6B6862"     # secondary text
FAINT = "9B9A94"    # captions / footer
DEFAULT_ACCENT = "1F6F43"  # forest green
GAP_RED = "B0432F"  # "no source" warning ink

DOC_FONT = "Aptos"


@dataclass
class ExportCitation:
    document_filename: str
    page: int | None = None


@dataclass
class ExportQuestion:
    requirement_id: str | None
    question_text: str
    answer: str
    citations: list[ExportCitation] = field(default_factory=list)
    gap_flag: str | None = None  # "ok" | "partial" | "no_source" | None


@dataclass
class ProposalSection:
    heading: str
    content: str  # "__QA_BLOCK__" expands into the full Q&A block in place


@dataclass
class DocGroup:
    heading: str
    items: list[ExportQuestion]


@dataclass
class ExportOptions:
    deal_name: str
    client_name: str | None = None
    org_name: str | None = None
    citation_style: str = "inline"  # "inline" | "footnote"
    sections: list[DocGroup] | None = None
    proposal_sections: list[ProposalSection] | None = None
    accent_color: str = DEFAULT_ACCENT


def _fmt_date() -> str:
    # Long human date, e.g. "July 13, 2026". Day formatted manually because
    # %-d isn't portable to Windows.
    d = date.today()
    return f"{d.strftime('%B')} {d.day}, {d.year}"


def _inline_citation_text(q: ExportQuestion) -> str:
    parts = [
        f"[Source: {c.document_filename}{f', p.{c.page}' if c.page is not None else ''}]"
        for c in q.citations
    ]
    return " ".join(parts)


def _groups(questions: list[ExportQuestion], opts: ExportOptions) -> list[DocGroup | None]:
    if opts.sections:
        return list(opts.sections)
    return [None]  # single untitled group over the flat question list


# =====================================================================
# DOCX
# =====================================================================

def _set_run(run, *, size_pt: float, color: str, bold=False, italic=False, superscript=False):
    run.font.name = DOC_FONT
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic
    if superscript:
        run.font.superscript = True
    # Ensure east-asian/complex-script also use the same face.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), DOC_FONT)


def _accent_rule(paragraph, color: str):
    """Draw a heavy bottom border under an (empty) paragraph — the cover rule."""
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "12", qn("w:space"): "1", qn("w:color"): color},
    )
    pbdr.append(bottom)
    ppr.append(pbdr)


def _footer_with_page_numbers(section, deal_name: str):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_text(text):
        r = p.add_run(text)
        _set_run(r, size_pt=9, color=FAINT)
        return r

    def add_field(instr):
        r = p.add_run()
        _set_run(r, size_pt=9, color=FAINT)
        fld_begin = r._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        r._element.append(fld_begin)
        r2 = p.add_run()
        _set_run(r2, size_pt=9, color=FAINT)
        instr_el = r2._element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
        instr_el.text = f" {instr} "
        r2._element.append(instr_el)
        r3 = p.add_run()
        _set_run(r3, size_pt=9, color=FAINT)
        fld_end = r3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        r3._element.append(fld_end)

    add_text(f"{deal_name} — Page ")
    add_field("PAGE")
    add_text(" of ")
    add_field("NUMPAGES")


def render_docx(questions: list[ExportQuestion], opts: ExportOptions) -> bytes:
    doc = Document()
    doc.core_properties.title = f"RFP Response — {opts.deal_name}"
    doc.core_properties.author = opts.org_name or "Klovered"
    doc.core_properties.comments = "Generated by Klovered"

    # Document-wide default: one clean sans, 11pt, comfortable ~1.3 line height.
    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.3
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), DOC_FONT)

    accent = opts.accent_color or DEFAULT_ACCENT

    section = doc.sections[0]
    _footer_with_page_numbers(section, opts.deal_name)

    # ---- cover block ----
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    _set_run(kicker.add_run("R F P   R E S P O N S E"), size_pt=9, color=accent, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_run(title.add_run(opts.deal_name), size_pt=23, color=INK, bold=True)

    if opts.client_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _set_run(p.add_run("Prepared for "), size_pt=11, color=FAINT)
        _set_run(p.add_run(opts.client_name), size_pt=11, color=MUTE)
    if opts.org_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _set_run(p.add_run("Submitted by "), size_pt=11, color=FAINT)
        _set_run(p.add_run(opts.org_name), size_pt=11, color=MUTE)

    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(10)
    _set_run(dp.add_run(_fmt_date()), size_pt=10, color=FAINT)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(24)
    _accent_rule(rule, accent)

    # ---- body ----
    groups = _groups(questions, opts)

    def render_qa_block(footnote_state: dict):
        for group in groups:
            if group is not None and group.heading:
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(10)
                _set_run(h.add_run(group.heading), size_pt=16, color=accent, bold=True)
            items = group.items if group is not None else questions
            for q in items:
                footnote_state["sources"].append(q)
                if q.requirement_id:
                    rp = doc.add_paragraph()
                    rp.paragraph_format.space_before = Pt(12)
                    rp.paragraph_format.space_after = Pt(3)
                    _set_run(rp.add_run(q.requirement_id), size_pt=10, color=DEFAULT_ACCENT, bold=True)
                qp = doc.add_paragraph()
                qp.paragraph_format.space_after = Pt(6)
                _set_run(qp.add_run(q.question_text), size_pt=12, color=INK, bold=True)

                if q.gap_flag == "no_source":
                    gp = doc.add_paragraph()
                    gp.paragraph_format.space_after = Pt(12)
                    _set_run(
                        gp.add_run(
                            "No source found in the knowledge base. This requirement "
                            "requires human review before submission."
                        ),
                        size_pt=11, color=GAP_RED, italic=True,
                    )
                    continue

                ap = doc.add_paragraph()
                ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                ap.paragraph_format.space_after = Pt(12)
                _set_run(ap.add_run(q.answer or "(no response)"), size_pt=11, color=INK)
                if opts.citation_style == "inline":
                    inline = _inline_citation_text(q)
                    if inline:
                        _set_run(ap.add_run(f" {inline}"), size_pt=10, color=MUTE)
                elif opts.citation_style == "footnote" and q.citations:
                    refs = []
                    for _c in q.citations:
                        footnote_state["counter"] += 1
                        refs.append(str(footnote_state["counter"]))
                    _set_run(
                        ap.add_run(f" [{', '.join(refs)}]"),
                        size_pt=9, color=DEFAULT_ACCENT, superscript=True,
                    )

    footnote_state = {"counter": 0, "sources": []}
    qa_rendered = False

    if opts.proposal_sections:
        for ps in opts.proposal_sections:
            if ps.content == "__QA_BLOCK__":
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(24)
                h.paragraph_format.space_after = Pt(10)
                _set_run(h.add_run(ps.heading), size_pt=16, color=accent, bold=True)
                render_qa_block(footnote_state)
                qa_rendered = True
                continue
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(8)
            _set_run(h.add_run(ps.heading), size_pt=16, color=accent, bold=True)
            for line in ps.content.split("\n"):
                if not line.strip():
                    continue
                lp = doc.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                lp.paragraph_format.space_after = Pt(8)
                _set_run(lp.add_run(line), size_pt=11, color=INK)

    if not qa_rendered:
        render_qa_block(footnote_state)

    if opts.citation_style == "footnote":
        rh = doc.add_paragraph()
        rh.paragraph_format.space_before = Pt(24)
        rh.paragraph_format.space_after = Pt(6)
        _set_run(rh.add_run("References"), size_pt=14, color=INK, bold=True)
        n = 0
        for q in footnote_state["sources"]:
            for c in q.citations:
                n += 1
                rp = doc.add_paragraph()
                num = rp.add_run(f"{n}. ")
                _set_run(num, size_pt=10, color=INK, bold=True)
                tail = f"{c.document_filename}{f' — page {c.page}' if c.page is not None else ''}"
                _set_run(rp.add_run(tail), size_pt=10, color=INK)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# =====================================================================
# PDF
# =====================================================================

def render_pdf(questions: list[ExportQuestion], opts: ExportOptions) -> bytes:
    # Imported lazily so the DOCX path (and tests that only touch it) don't pay
    # reportlab's import cost.
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        BaseDocTemplate,
        Flowable,
        Frame,
        HRFlowable,
        PageTemplate,
        Paragraph,
        Spacer,
    )
    from reportlab.lib.styles import ParagraphStyle
    from xml.sax.saxutils import escape as xml_escape

    accent = HexColor(f"#{(opts.accent_color or DEFAULT_ACCENT).lstrip('#')}")
    ink = HexColor(f"#{INK}")
    mute = HexColor(f"#{MUTE}")
    faint = HexColor(f"#{FAINT}")
    gap = HexColor(f"#{GAP_RED}")

    def style(name, **kw):
        base = dict(fontName="Helvetica", textColor=ink, fontSize=10.5, leading=15)
        base.update(kw)
        return ParagraphStyle(name, **base)

    st_kicker = style("kicker", fontName="Helvetica-Bold", fontSize=9, textColor=accent, leading=12)
    st_title = style("title", fontName="Helvetica-Bold", fontSize=24, textColor=ink, leading=28)
    st_meta = style("meta", fontSize=11, textColor=mute, leading=15)
    st_meta_faint = style("metafaint", fontSize=11, textColor=faint, leading=15)
    st_date = style("date", fontSize=10, textColor=faint, leading=13)
    st_group = style("group", fontName="Helvetica-Bold", fontSize=16, textColor=accent, leading=20)
    st_section = style("section", fontName="Helvetica-Bold", fontSize=15, textColor=accent, leading=19)
    st_req = style("req", fontName="Helvetica-Bold", fontSize=9.5, textColor=accent, leading=12)
    st_question = style("question", fontName="Helvetica-Bold", fontSize=12.5, textColor=ink, leading=16)
    st_answer = style("answer", fontSize=10.5, textColor=ink, leading=15.5, alignment=TA_JUSTIFY)
    st_gap = style("gap", fontSize=10.5, textColor=gap, leading=15, fontName="Helvetica-Oblique")
    st_cite = style("cite", fontSize=9, textColor=mute, leading=12)
    st_body = style("body", fontSize=10.5, textColor=ink, leading=15.5, alignment=TA_JUSTIFY)
    st_refh = style("refh", fontName="Helvetica-Bold", fontSize=14, textColor=ink, leading=18)
    st_ref = style("ref", fontSize=10, textColor=ink, leading=14)

    def esc(s: str) -> str:
        return xml_escape(s or "")

    story: list[Flowable] = []
    story.append(Paragraph("R F P &#160;&#160; R E S P O N S E", st_kicker))
    story.append(Spacer(1, 4))
    story.append(Paragraph(esc(opts.deal_name), st_title))
    story.append(Spacer(1, 8))
    if opts.client_name:
        story.append(
            Paragraph(f'<font color="#{FAINT}">Prepared for </font>{esc(opts.client_name)}', st_meta)
        )
    if opts.org_name:
        story.append(
            Paragraph(f'<font color="#{FAINT}">Submitted by </font>{esc(opts.org_name)}', st_meta)
        )
    story.append(Spacer(1, 6))
    story.append(Paragraph(_fmt_date(), st_date))
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=1.5, color=accent, spaceAfter=20))

    groups = _groups(questions, opts)

    def qa_flowables(footnote_state: dict) -> list[Flowable]:
        out: list[Flowable] = []
        for gi, group in enumerate(groups):
            if group is not None and group.heading:
                if gi > 0:
                    out.append(Spacer(1, 12))
                out.append(Paragraph(esc(group.heading), st_group))
                out.append(Spacer(1, 8))
            items = group.items if group is not None else questions
            for q in items:
                footnote_state["sources"].append(q)
                if q.requirement_id:
                    out.append(Spacer(1, 8))
                    out.append(Paragraph(esc(q.requirement_id), st_req))
                    out.append(Spacer(1, 2))
                else:
                    out.append(Spacer(1, 8))
                out.append(Paragraph(esc(q.question_text), st_question))
                out.append(Spacer(1, 5))
                if q.gap_flag == "no_source":
                    out.append(Paragraph(
                        "No source found in the knowledge base. Human review required "
                        "before submission.", st_gap,
                    ))
                    continue
                out.append(Paragraph(esc(q.answer or "(no response)"), st_answer))
                if opts.citation_style == "inline":
                    inline = _inline_citation_text(q)
                    if inline:
                        out.append(Spacer(1, 2))
                        out.append(Paragraph(esc(inline), st_cite))
                elif opts.citation_style == "footnote" and q.citations:
                    refs = []
                    for _c in q.citations:
                        footnote_state["counter"] += 1
                        refs.append(str(footnote_state["counter"]))
                    out.append(Spacer(1, 2))
                    out.append(Paragraph(f'<super>[{", ".join(refs)}]</super>', st_cite))
        return out

    footnote_state = {"counter": 0, "sources": []}
    qa_rendered = False

    if opts.proposal_sections:
        for ps in opts.proposal_sections:
            story.append(Spacer(1, 16))
            story.append(Paragraph(esc(ps.heading), st_section))
            story.append(Spacer(1, 8))
            if ps.content == "__QA_BLOCK__":
                story.extend(qa_flowables(footnote_state))
                qa_rendered = True
                continue
            for line in ps.content.split("\n"):
                if line.strip():
                    story.append(Paragraph(esc(line), st_body))
                    story.append(Spacer(1, 6))

    if not qa_rendered:
        story.extend(qa_flowables(footnote_state))

    if opts.citation_style == "footnote":
        story.append(Spacer(1, 20))
        story.append(Paragraph("References", st_refh))
        story.append(Spacer(1, 6))
        n = 0
        for q in footnote_state["sources"]:
            for c in q.citations:
                n += 1
                tail = f"{c.document_filename}{f' — page {c.page}' if c.page is not None else ''}"
                story.append(Paragraph(f"<b>{n}.</b> {esc(tail)}", st_ref))

    buf = io.BytesIO()
    doc_name = opts.deal_name

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(faint)
        text = f"{doc_name} — Page {doc_.page}"
        canvas.drawCentredString(LETTER[0] / 2, 0.5 * inch, text)
        canvas.restoreState()

    doc = BaseDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.85 * inch, rightMargin=0.85 * inch,
        topMargin=0.85 * inch, bottomMargin=0.85 * inch,
        title=f"RFP Response — {opts.deal_name}", author=opts.org_name or "Klovered",
    )
    frame = Frame(
        doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="body",
        leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0,
    )
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=footer)])
    doc.build(story)
    return buf.getvalue()
