"""Branded .docx export with inline citations or footnotes.

Python port of ``lib/docx-export.ts`` using python-docx. Produces a real Word
file that opens in MS Word and Google Docs.
"""

from __future__ import annotations

import datetime as _dt
import io
from dataclasses import dataclass, field
from typing import Any, Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Document typeface. Geist (the product UI font) isn't installed in Word or
# Google Docs, so we set the closest widely-available modern sans as the
# default and let the reader substitute gracefully.
DOC_FONT = "Aptos"
INK = "1A1A17"  # near-black body ink (matches the app's --fg ramp)
MUTE = "6B6862"  # secondary text
FAINT = "9B9A94"  # captions / footer


# ---------------------------------------------------------------------------
# Types
# ---------------------------------------------------------------------------

GapFlag = Literal["ok", "partial", "no_source"]
CitationStyle = Literal["inline", "footnote"]


@dataclass
class ExportQuestion:
    requirement_id: str | None
    question_text: str
    citations: list[dict[str, Any]]  # {document_filename: str, page: int|None}
    answer: str
    gap_flag: GapFlag | None


@dataclass
class ExportOptions:
    deal_name: str
    client_name: str | None
    org_name: str | None
    citation_style: CitationStyle
    sections: list[dict[str, Any]] | None = None  # {heading, items: list[ExportQuestion]}
    proposal_sections: list[dict[str, Any]] | None = None  # {heading, content}
    accent_color: str | None = None
    logo: dict[str, Any] | None = None  # {data: bytes, ext: "png"|"jpg"}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _add_run(paragraph, text: str, *, bold=False, italics=False, size=None,
             color=None, superscript=False, font=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italics
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if superscript:
        run.font.superscript = True
    if font is not None:
        run.font.name = font
    return run


def _half_pt_to_pt(half_points: int) -> float:
    return half_points / 2.0


def _add_field(paragraph, field_code: str):
    """Insert a Word field code (e.g. PAGE, NUMPAGES) into a paragraph."""
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    return run


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def render_docx(questions: list[ExportQuestion], opts: ExportOptions) -> bytes:
    document = Document()

    # Document-wide typographic defaults.
    normal = document.styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(_half_pt_to_pt(22))
    normal.font.color.rgb = RGBColor.from_string(INK)

    accent_cover = opts.accent_color or "1F6F43"
    accent = opts.accent_color or "1F6F43"

    # ---- Cover / title block ----
    if opts.logo:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(io.BytesIO(opts.logo["data"]), width=Pt(105), height=Pt(45))
        p.paragraph_format.space_after = Pt(12)

    kicker = document.add_paragraph()
    _add_run(kicker, "R F P   R E S P O N S E", bold=True, size=_half_pt_to_pt(18), color=accent_cover)
    kicker.paragraph_format.space_after = Pt(3)

    title = document.add_paragraph()
    _add_run(title, opts.deal_name, bold=True, size=_half_pt_to_pt(46), color=INK)
    title.paragraph_format.space_after = Pt(4)

    if opts.client_name:
        p = document.add_paragraph()
        _add_run(p, "Prepared for ", size=_half_pt_to_pt(22), color=FAINT)
        _add_run(p, opts.client_name, size=_half_pt_to_pt(22), color=MUTE)
        p.paragraph_format.space_after = Pt(1)

    if opts.org_name:
        p = document.add_paragraph()
        _add_run(p, "Submitted by ", size=_half_pt_to_pt(22), color=FAINT)
        _add_run(p, opts.org_name, size=_half_pt_to_pt(22), color=MUTE)
        p.paragraph_format.space_after = Pt(1)

    date_p = document.add_paragraph()
    _today = _dt.date.today()
    _add_run(
        date_p,
        f"{_today:%B} {_today.day}, {_today:%Y}",
        size=_half_pt_to_pt(20),
        color=FAINT,
    )
    date_p.paragraph_format.space_after = Pt(10)

    # Accent divider rule under the cover block.
    divider = document.add_paragraph()
    divider.paragraph_format.space_after = Pt(24)
    p_pr = divider._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), accent_cover)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    # ---- Q&A block builder ----
    footnote_counter = 0
    footnote_source_list: list[ExportQuestion] = []

    groups: list[dict[str, Any]]
    if opts.sections:
        groups = [{"heading": s.get("heading"), "items": s["items"]} for s in opts.sections]
    else:
        groups = [{"heading": None, "items": questions}]

    def render_qa_block():
        nonlocal footnote_counter
        for group in groups:
            if group.get("heading"):
                h = document.add_paragraph()
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(10)
                _add_run(h, group["heading"], bold=True, size=_half_pt_to_pt(32), color=accent)

            for q in group["items"]:
                footnote_source_list.append(q)
                if q.requirement_id:
                    rp = document.add_paragraph()
                    rp.paragraph_format.space_before = Pt(12)
                    rp.paragraph_format.space_after = Pt(3)
                    _add_run(rp, q.requirement_id, bold=True, size=_half_pt_to_pt(20), color="1F6F43")

                qp = document.add_paragraph()
                qp.paragraph_format.space_after = Pt(6)
                _add_run(qp, q.question_text, bold=True, size=_half_pt_to_pt(24), color="1A1A17")

                if q.gap_flag == "no_source":
                    gp = document.add_paragraph()
                    gp.paragraph_format.space_after = Pt(12)
                    _add_run(
                        gp,
                        "No source found in the knowledge base. This requirement requires "
                        "human review before submission.",
                        italics=True,
                        color="B0432F",
                        size=_half_pt_to_pt(22),
                    )
                    continue

                ap = document.add_paragraph()
                ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                ap.paragraph_format.space_after = Pt(12)
                _add_run(ap, q.answer or "(no response)", size=_half_pt_to_pt(22))

                if opts.citation_style == "inline":
                    inline = " ".join(
                        f"[Source: {c['document_filename']}"
                        + (f", p.{c['page']}" if c.get("page") is not None else "")
                        + "]"
                        for c in q.citations
                    )
                    if inline:
                        _add_run(ap, f" {inline}", size=_half_pt_to_pt(20), color="6B6862")
                elif opts.citation_style == "footnote" and q.citations:
                    refs = []
                    for _c in q.citations:
                        footnote_counter += 1
                        refs.append(str(footnote_counter))
                    _add_run(
                        ap,
                        f" [{', '.join(refs)}]",
                        superscript=True,
                        size=_half_pt_to_pt(18),
                        color="1F6F43",
                    )

    qa_rendered = False

    if opts.proposal_sections:
        for ps in opts.proposal_sections:
            if ps.get("content") == "__QA_BLOCK__":
                h = document.add_paragraph()
                h.paragraph_format.space_before = Pt(24)
                h.paragraph_format.space_after = Pt(10)
                _add_run(h, ps["heading"], bold=True, size=_half_pt_to_pt(32), color=accent)
                render_qa_block()
                qa_rendered = True
                continue

            h = document.add_paragraph()
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(8)
            _add_run(h, ps["heading"], bold=True, size=_half_pt_to_pt(32), color=accent)

            lines = [l for l in ps.get("content", "").split("\n") if l.strip() != ""]
            for line in lines:
                lp = document.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                lp.paragraph_format.space_after = Pt(8)
                _add_run(lp, line, size=_half_pt_to_pt(22))

    if not qa_rendered:
        render_qa_block()

    if opts.citation_style == "footnote":
        h = document.add_paragraph()
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(6)
        _add_run(h, "References", bold=True, size=_half_pt_to_pt(28))

        n = 0
        for q in footnote_source_list:
            for c in q.citations:
                n += 1
                rp = document.add_paragraph()
                _add_run(rp, f"{n}. ", bold=True, size=_half_pt_to_pt(20))
                page_suffix = f" — page {c['page']}" if c.get("page") is not None else ""
                _add_run(rp, f"{c['document_filename']}{page_suffix}", size=_half_pt_to_pt(20))

    # ---- Footer with page numbers ----
    section = document.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(footer_p, f"{opts.deal_name} — Page ", size=_half_pt_to_pt(18), color="9B9A94")
    _add_field(footer_p, "PAGE")
    _add_run(footer_p, " of ", size=_half_pt_to_pt(18), color="9B9A94")
    _add_field(footer_p, "NUMPAGES")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
