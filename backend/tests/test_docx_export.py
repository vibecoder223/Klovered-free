import io

from docx import Document

from app.pipeline.docx_export import ExportOptions, ExportQuestion, render_docx


def _sample_questions():
    return [
        ExportQuestion(
            requirement_id="REQ-1",
            question_text="Do you support SSO?",
            citations=[{"document_filename": "security.pdf", "page": 3}],
            answer="Yes, we support SAML and OIDC based SSO.",
            gap_flag="ok",
        ),
        ExportQuestion(
            requirement_id="REQ-2",
            question_text="Do you support FedRAMP?",
            citations=[],
            answer="",
            gap_flag="no_source",
        ),
    ]


def _all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def test_render_docx_inline_citation_style():
    questions = _sample_questions()
    opts = ExportOptions(
        deal_name="Acme Deal",
        client_name="Acme Corp",
        org_name="Klovered",
        citation_style="inline",
    )
    data = render_docx(questions, opts)
    doc = Document(io.BytesIO(data))
    text = _all_text(doc)

    assert "Acme Deal" in text
    assert "Do you support SSO?" in text
    assert "Do you support FedRAMP?" in text
    assert "Yes, we support SAML and OIDC based SSO." in text
    assert "[Source: security.pdf, p.3]" in text
    assert "No source found in the knowledge base. This requirement requires human review before submission." in text


def test_render_docx_footnote_citation_style():
    questions = _sample_questions()
    opts = ExportOptions(
        deal_name="Acme Deal",
        client_name=None,
        org_name=None,
        citation_style="footnote",
    )
    data = render_docx(questions, opts)
    doc = Document(io.BytesIO(data))
    text = _all_text(doc)

    assert any(p.text.strip() == "References" for p in doc.paragraphs)
    assert any("security.pdf" in p.text and "page 3" in p.text for p in doc.paragraphs)
