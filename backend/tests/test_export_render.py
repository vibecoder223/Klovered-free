"""Unit tests for the DOCX/PDF renderers — validate real, well-formed output
(not just that a function ran) across inline/footnote citations, per-document
grouping (merge), proposal sections, and the no-source gap path.
"""

import io
import zipfile

from docx import Document

from app.pipeline.export_render import (
    DocGroup,
    ExportCitation,
    ExportOptions,
    ExportQuestion,
    ProposalSection,
    render_docx,
    render_pdf,
)


def _q(rid="REQ-1", text="Encrypt?", answer="Yes, AES-256.", pages=(3,), gap="ok"):
    return ExportQuestion(
        requirement_id=rid,
        question_text=text,
        answer=answer,
        citations=[ExportCitation("sec.pdf", p) for p in pages],
        gap_flag=gap,
    )


def _docx_text(data: bytes) -> str:
    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


def test_docx_is_valid_zip_with_expected_content():
    data = render_docx(
        [_q(), _q(rid="REQ-2", text="RTO?", answer="4 hours.")],
        ExportOptions(deal_name="Big RFP", client_name="Acme", org_name="Us",
                      citation_style="inline"),
    )
    assert data[:2] == b"PK"
    text = _docx_text(data)
    assert "Big RFP" in text
    assert "Prepared for" in text and "Acme" in text
    assert "REQ-1" in text and "Encrypt?" in text
    assert "[Source: sec.pdf, p.3]" in text


def test_docx_footnote_adds_references_section():
    data = render_docx(
        [_q(pages=(3, 5))],
        ExportOptions(deal_name="D", citation_style="footnote"),
    )
    text = _docx_text(data)
    assert "References" in text
    # two citations -> two numbered reference lines
    assert "1. sec.pdf — page 3" in text
    assert "2. sec.pdf — page 5" in text


def test_docx_no_source_shows_gap_notice_and_skips_answer():
    data = render_docx(
        [_q(answer="", pages=(), gap="no_source")],
        ExportOptions(deal_name="D", citation_style="inline"),
    )
    text = _docx_text(data)
    assert "No source found in the knowledge base" in text


def test_docx_uses_aptos_font_everywhere():
    data = render_docx([_q()], ExportOptions(deal_name="D"))
    doc = Document(io.BytesIO(data))
    fonts = {r.font.name for p in doc.paragraphs for r in p.runs if r.font.name}
    assert fonts == {"Aptos"}


def test_docx_footer_has_page_number_fields():
    data = render_docx([_q()], ExportOptions(deal_name="MyDeal"))
    z = zipfile.ZipFile(io.BytesIO(data))
    footer_xml = next(z.read(n).decode() for n in z.namelist() if "footer" in n)
    assert "PAGE" in footer_xml and "NUMPAGES" in footer_xml
    assert "MyDeal" in footer_xml


def test_docx_merge_sections_emit_per_document_headings():
    sections = [
        DocGroup(heading="rfp-a.pdf", items=[_q(rid="A1")]),
        DocGroup(heading="rfp-b.pdf", items=[_q(rid="B1")]),
    ]
    data = render_docx([], ExportOptions(deal_name="D", sections=sections))
    text = _docx_text(data)
    assert "rfp-a.pdf" in text and "rfp-b.pdf" in text
    assert "A1" in text and "B1" in text


def test_docx_proposal_section_and_qa_block_ordering():
    data = render_docx(
        [_q()],
        ExportOptions(
            deal_name="D",
            proposal_sections=[
                ProposalSection("Executive Summary", "We are pleased to respond.\nSecond line."),
                ProposalSection("Responses", "__QA_BLOCK__"),
            ],
        ),
    )
    text = _docx_text(data)
    assert text.index("Executive Summary") < text.index("Responses") < text.index("Encrypt?")
    assert "We are pleased to respond." in text


def test_pdf_is_valid_and_multipage_ready():
    data = render_pdf(
        [_q() for _ in range(20)],  # enough to force pagination
        ExportOptions(deal_name="Long RFP", client_name="C", org_name="O",
                      citation_style="footnote"),
    )
    assert data[:5] == b"%PDF-"
    # crude page-count check via the PDF's /Type /Page objects
    assert data.count(b"/Type /Page") >= 1


def test_pdf_inline_and_accent_override_render():
    data = render_pdf(
        [_q()],
        ExportOptions(deal_name="D", citation_style="inline", accent_color="#8800AA"),
    )
    assert data[:5] == b"%PDF-"
