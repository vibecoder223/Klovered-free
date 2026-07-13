import io

import fitz
import pytest
from docx import Document

from app.pipeline.parse import parse_document


def test_txt_two_paragraphs():
    doc = parse_document(b"Para one.\n\nPara two.", "text/plain", "x.txt")
    assert doc.page_count == 1
    paras = [b for b in doc.blocks if b.type == "paragraph"]
    assert len(paras) == 2
    assert paras[0].text == "Para one."
    assert paras[1].text == "Para two."


def test_docx_heading_paragraph_list():
    d = Document()
    d.add_heading("Section One", level=1)
    d.add_paragraph("This is a normal paragraph of text.")
    d.add_paragraph("First bullet item", style="List Bullet")
    buf = io.BytesIO()
    d.save(buf)
    data = buf.getvalue()

    parsed = parse_document(data, None, "test.docx")
    assert parsed.page_count == 1

    headings = [b for b in parsed.blocks if b.type == "heading"]
    paragraphs = [b for b in parsed.blocks if b.type == "paragraph"]
    assert any("Section One" in h.text for h in headings)
    assert any("normal paragraph" in p.text for p in paragraphs)
    # Bullet may or may not come through as <li> depending on mammoth's style
    # mapping; only assert on the heading + paragraph which are guaranteed.


def test_pdf_single_page():
    d = fitz.open()
    page = d.new_page()
    page.insert_text((72, 72), "Hello requirement text")
    data = d.tobytes()
    d.close()

    parsed = parse_document(data, "application/pdf", "test.pdf")
    assert parsed.page_count == 1
    assert "Hello" in parsed.raw_text


def test_unsupported_type_raises():
    with pytest.raises(ValueError):
        parse_document(b"whatever", "application/zip", "x.zip")
